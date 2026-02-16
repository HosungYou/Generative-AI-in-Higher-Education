#!/usr/bin/env python3
"""Convert Paper A and Paper B markdown manuscripts to APA 7th edition Word documents."""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════════════════════════════════
# APA 7th STYLE CONSTANTS
# ═══════════════════════════════════════════════════════════════
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = 2.0  # Double-spaced
MARGIN = Inches(1)
FIRST_LINE_INDENT = Inches(0.5)
HANGING_INDENT = Inches(0.5)


def create_apa_document():
    """Create a blank document with APA 7th formatting defaults."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    # Default paragraph style
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = FIRST_LINE_INDENT

    return doc


def add_page_number(doc):
    """Add page numbers to footer (right-aligned)."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        fld_xml = (
            '<w:fldSimple {} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        ).format(nsdecls("w"))
        fld = parse_xml(fld_xml)
        run._r.append(fld)
        rpr = run._r.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>')
        rpr.append(rFonts)
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="24"/>')
        rpr.append(sz)


def add_running_head(doc, short_title):
    """Add running head to header."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(short_title.upper())
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE


def set_paragraph_apa(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=True,
                       space_before=0, space_after=0):
    """Apply APA formatting to a paragraph."""
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent:
        pf.first_line_indent = FIRST_LINE_INDENT
    else:
        pf.first_line_indent = Pt(0)


def add_run_with_format(paragraph, text, bold=False, italic=False):
    """Add a run with specific formatting."""
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    run.italic = italic
    return run


def parse_inline_formatting(paragraph, text):
    """Parse markdown inline formatting (**bold**, *italic*) and add runs."""
    # Pattern: **bold**, *italic* (handle bold first to avoid conflict)
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for match in pattern.finditer(text):
        # Add text before this match
        if match.start() > pos:
            add_run_with_format(paragraph, text[pos:match.start()])

        if match.group(2):  # ***bold italic***
            add_run_with_format(paragraph, match.group(2), bold=True, italic=True)
        elif match.group(3):  # **bold**
            add_run_with_format(paragraph, match.group(3), bold=True)
        elif match.group(4):  # *italic*
            add_run_with_format(paragraph, match.group(4), italic=True)

        pos = match.end()

    # Add remaining text
    if pos < len(text):
        add_run_with_format(paragraph, text[pos:])


def add_apa_heading(doc, text, level=1):
    """Add APA 7th heading.
    Level 1: Centered, Bold, Title Case
    Level 2: Left-aligned, Bold, Title Case
    Level 3: Left-aligned, Bold Italic, Title Case
    Level 4: Indented, Bold, Title Case, ending with period (inline start)
    Level 5: Indented, Bold Italic, Title Case, ending with period (inline start)
    """
    # Clean markdown formatting from heading text
    clean = re.sub(r'\*+', '', text).strip()

    if level <= 3:
        p = doc.add_paragraph()
        set_paragraph_apa(p, first_indent=False)

        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_with_format(p, clean, bold=True)
        elif level == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run_with_format(p, clean, bold=True)
        elif level == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run_with_format(p, clean, bold=True, italic=True)
        return p
    else:
        # Levels 4-5 are inline (returned for caller to append body text)
        p = doc.add_paragraph()
        set_paragraph_apa(p, first_indent=True)
        if level == 4:
            add_run_with_format(p, clean + ". ", bold=True)
        else:
            add_run_with_format(p, clean + ". ", bold=True, italic=True)
        return p


def add_apa_table(doc, header_row, data_rows, note_text=None):
    """Add an APA-style table (no vertical borders, horizontal rules only)."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove all borders first
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    def set_cell_border(cell, top=False, bottom=False):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        if top:
            tcBorders.append(parse_xml(
                f'<w:top {nsdecls("w")} w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            ))
        if bottom:
            tcBorders.append(parse_xml(
                f'<w:bottom {nsdecls("w")} w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            ))
        tcPr.append(tcBorders)

    # Header row
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.first_line_indent = Pt(0)
        run = p.add_run(h.strip())
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.bold = True
        # Top and bottom border on header
        set_cell_border(cell, top=True, bottom=True)

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        is_last = (r_idx == len(data_rows) - 1)
        for c_idx in range(num_cols):
            cell = table.rows[r_idx + 1].cells[c_idx]
            val = row_data[c_idx].strip() if c_idx < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            # Left-align first column, center others
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.first_line_indent = Pt(0)
            # Parse bold/italic in cell text
            clean_val = re.sub(r'\*+', '', val)
            is_bold = val.startswith("**") and val.endswith("**")
            is_italic = val.startswith("*") and val.endswith("*") and not is_bold
            run = p.add_run(clean_val)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.bold = is_bold
            run.italic = is_italic
            # Bottom border on last row
            if is_last:
                set_cell_border(cell, bottom=True)

    # Table note
    if note_text:
        note_p = doc.add_paragraph()
        set_paragraph_apa(note_p, first_indent=False)
        note_p.paragraph_format.space_before = Pt(4)
        clean_note = re.sub(r'\*+', '', note_text)
        run = note_p.add_run("Note. ")
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.italic = True
        run = note_p.add_run(clean_note)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)


def add_reference(doc, ref_text):
    """Add a reference entry with hanging indent."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(-0.5)
    pf.left_indent = Inches(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parse_inline_formatting(p, ref_text)


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (header, rows, note, end_idx)."""
    header_line = lines[start_idx].strip()
    cols = [c.strip() for c in header_line.split("|") if c.strip()]

    # Skip separator line
    sep_idx = start_idx + 1
    if sep_idx < len(lines) and re.match(r'\s*\|[-:|]+\|', lines[sep_idx]):
        sep_idx += 1
    else:
        return None, None, None, start_idx

    rows = []
    idx = sep_idx
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith("|"):
            break
        row_cols = [c.strip() for c in line.split("|") if c.strip() != ""]
        # Handle empty cells in split
        raw = line.strip("|").split("|")
        row_cols = [c.strip() for c in raw]
        rows.append(row_cols)
        idx += 1

    # Look for note after table
    note = None
    if idx < len(lines) and lines[idx].strip().startswith("*Note"):
        note_text = lines[idx].strip()
        # Clean leading *Note.* format
        note_text = re.sub(r'^\*Note\.\*\s*', '', note_text)
        note = note_text
        idx += 1

    return cols, rows, note, idx


def process_markdown_to_apa(md_path, doc, metadata):
    """Process a markdown file and populate the APA document."""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Remove HTML comments
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)

    lines = content.split("\n")
    in_references = False
    ref_buffer = []
    skip_until = -1
    i = 0

    # ── Title Page ──
    doc.add_paragraph()  # Blank line at top
    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    tp = doc.add_paragraph()
    set_paragraph_apa(tp, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_run_with_format(tp, metadata["title"], bold=True)

    # Blank line
    doc.add_paragraph()

    # Author
    ap = doc.add_paragraph()
    set_paragraph_apa(ap, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_run_with_format(ap, metadata["author"])

    # Affiliation
    aff = doc.add_paragraph()
    set_paragraph_apa(aff, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_run_with_format(aff, metadata["affiliation"])

    # Author Note
    doc.add_paragraph()
    doc.add_paragraph()
    an_head = doc.add_paragraph()
    set_paragraph_apa(an_head, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_run_with_format(an_head, "Author Note", bold=True)

    for note_line in metadata.get("author_note", []):
        anp = doc.add_paragraph()
        set_paragraph_apa(anp, first_indent=True)
        parse_inline_formatting(anp, note_line)

    # Page break before abstract
    doc.add_page_break()

    # ── Abstract ──
    abs_head = doc.add_paragraph()
    set_paragraph_apa(abs_head, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_run_with_format(abs_head, "Abstract", bold=True)

    # Find abstract in content
    abstract_text = ""
    keywords_text = ""
    abs_start = content.find("## Abstract")
    if abs_start >= 0:
        abs_section = content[abs_start:]
        # Find the abstract body (between ## Abstract and next ##)
        abs_match = re.search(r'## Abstract\s*\n\n(.+?)\n\n\*Keywords:\*\s*(.+?)\n', abs_section, re.DOTALL)
        if abs_match:
            abstract_text = abs_match.group(1).strip()
            keywords_text = abs_match.group(2).strip()

    if abstract_text:
        abs_p = doc.add_paragraph()
        set_paragraph_apa(abs_p, first_indent=False)
        parse_inline_formatting(abs_p, abstract_text)

    if keywords_text:
        kw_p = doc.add_paragraph()
        set_paragraph_apa(kw_p, first_indent=True)
        add_run_with_format(kw_p, "Keywords: ", italic=True)
        add_run_with_format(kw_p, keywords_text)

    # Page break before body
    doc.add_page_break()

    # Title again at start of body
    body_title = doc.add_paragraph()
    set_paragraph_apa(body_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_run_with_format(body_title, metadata["title"], bold=True)

    # ── Process body ──
    # Skip to after Abstract and Keywords
    body_start = content.find("---", content.find("*Keywords:*"))
    if body_start < 0:
        body_start = content.find("## ", content.find("*Keywords:*"))

    # Find the first heading after abstract/keywords that's the repeated title or intro
    body_content = content[body_start:] if body_start > 0 else content

    lines = body_content.split("\n")
    i = 0

    while i < len(lines):
        if i < skip_until:
            i += 1
            continue

        line = lines[i]
        stripped = line.strip()

        # Skip empty lines, horizontal rules, repeated title
        if stripped == "" or stripped == "---":
            i += 1
            continue

        # Skip the repeated title heading (matches main title)
        if stripped.startswith("## ") and metadata["title_short"] in stripped:
            i += 1
            continue

        # Skip version line and author info that's already on title page
        if stripped.startswith("**Version:") or stripped.startswith("**Hosung You**") or stripped == "College of Education, Pennsylvania State University":
            i += 1
            continue

        # Skip author note block (already handled)
        if stripped.startswith("**Author Note**"):
            # Skip until next ---
            i += 1
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # Skip abstract section (already handled)
        if stripped == "## Abstract":
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("## "):
                i += 1
            continue

        # ── References section ──
        if stripped == "## References":
            doc.add_page_break()
            add_apa_heading(doc, "References", level=1)
            in_references = True
            i += 1
            continue

        if in_references:
            if stripped.startswith("## ") or stripped.startswith("# "):
                # End of references
                in_references = False
                if ref_buffer:
                    add_reference(doc, " ".join(ref_buffer))
                    ref_buffer = []
                # Don't increment, let heading be processed
                continue

            if stripped == "":
                # Empty line = end of a reference entry
                if ref_buffer:
                    add_reference(doc, " ".join(ref_buffer))
                    ref_buffer = []
            else:
                ref_buffer.append(stripped)
            i += 1
            continue

        # ── Table handling ──
        if stripped.startswith("**Table ") and not stripped.startswith("**Table of"):
            # Table title (e.g., **Table 3**)
            table_num = re.sub(r'\*+', '', stripped).strip()
            tp = doc.add_paragraph()
            set_paragraph_apa(tp, first_indent=False, space_before=12)
            add_run_with_format(tp, table_num, bold=True)
            i += 1
            # Next line should be italic table description
            if i < len(lines):
                desc = lines[i].strip()
                if desc.startswith("*") and desc.endswith("*"):
                    desc_clean = desc.strip("*").strip()
                    td = doc.add_paragraph()
                    set_paragraph_apa(td, first_indent=False)
                    add_run_with_format(td, desc_clean, italic=True)
                    i += 1
            # Skip blank line before table
            while i < len(lines) and lines[i].strip() == "":
                i += 1
            # Parse the actual table
            if i < len(lines) and lines[i].strip().startswith("|"):
                header, rows, note, end_idx = parse_markdown_table(lines, i)
                if header and rows:
                    add_apa_table(doc, header, rows, note)
                i = end_idx
            continue

        # Standalone table (no preceding title)
        if stripped.startswith("|") and not in_references:
            # Check if previous non-empty line was a table title
            header, rows, note, end_idx = parse_markdown_table(lines, i)
            if header and rows:
                add_apa_table(doc, header, rows, note)
            i = end_idx
            continue

        # ── Headings ──
        if stripped.startswith("#"):
            heading_match = re.match(r'^(#{1,5})\s+(.+)', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()

                # Map markdown heading levels to APA levels
                # ## = Level 1, ### = Level 2, #### = Level 3
                apa_level = level - 1 if level >= 2 else 1

                if heading_text == "References":
                    doc.add_page_break()

                add_apa_heading(doc, heading_text, level=apa_level)
                i += 1
                continue

        # ── Figure placeholders ──
        if stripped.startswith("[Insert Figure"):
            fp = doc.add_paragraph()
            set_paragraph_apa(fp, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False,
                             space_before=12, space_after=12)
            add_run_with_format(fp, stripped.strip("[]"), italic=True)
            i += 1
            continue

        # ── Bulleted lists ──
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.25)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Add bullet character
            parse_inline_formatting(p, text)
            # Prepend bullet
            first_run = p.runs[0] if p.runs else p.add_run("")
            first_run.text = "\u2022 " + first_run.text
            i += 1
            continue

        # ── Numbered lists ──
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.25)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parse_inline_formatting(p, text)
            first_run = p.runs[0] if p.runs else p.add_run("")
            first_run.text = f"{num}. " + first_run.text
            i += 1
            continue

        # ── Regular paragraph ──
        if stripped and not stripped.startswith("|"):
            p = doc.add_paragraph()
            set_paragraph_apa(p, first_indent=True)
            parse_inline_formatting(p, stripped)
            i += 1
            continue

        i += 1

    # Flush remaining references
    if ref_buffer:
        add_reference(doc, " ".join(ref_buffer))


# ═══════════════════════════════════════════════════════════════
# PAPER A
# ═══════════════════════════════════════════════════════════════
def convert_paper_a():
    base = "/Users/hosung/Generative-AI-in-Higher-Education"
    md_path = os.path.join(base, "manuscript/versions/Paper_A_MetaAnalysis_v4.0.md")
    out_path = os.path.join(base, "manuscript/versions/Paper_A_MetaAnalysis_v4.0.docx")

    metadata = {
        "title": "Generative AI in Higher Education: A Three-Level Meta-Analysis "
                 "Revealing Cognitive Dependency in Metacognitive Outcomes",
        "title_short": "Generative AI in Higher Education",
        "running_head": "GENERATIVE AI IN HIGHER EDUCATION",
        "author": "Hosung You",
        "affiliation": "College of Education, Pennsylvania State University",
        "author_note": [
            "Hosung You  https://orcid.org/[ORCID-ID]",
            "Correspondence concerning this article should be addressed to "
            "Hosung You, College of Education, Pennsylvania State University, "
            "University Park, PA 16802. Email: hosung@psu.edu",
            "Data Availability Statement: The dataset, analysis code, and "
            "supplementary materials are available at [OSF Repository Link].",
            "Conflict of Interest: The author declares no conflicts of interest.",
            "Funding: This research received no external funding.",
        ],
    }

    doc = create_apa_document()
    add_running_head(doc, metadata["running_head"])
    add_page_number(doc)
    process_markdown_to_apa(md_path, doc, metadata)

    doc.save(out_path)
    size_kb = os.path.getsize(out_path) / 1024
    print(f"  Paper A: {out_path}")
    print(f"  Size: {size_kb:.1f} KB")
    return out_path


# ═══════════════════════════════════════════════════════════════
# PAPER B
# ═══════════════════════════════════════════════════════════════
def convert_paper_b():
    base = "/Users/hosung/Generative-AI-in-Higher-Education"
    md_path = os.path.join(base, "manuscript/versions/Paper_B_AI_Coding_Methodology_v1.0.md")
    out_path = os.path.join(base, "manuscript/versions/Paper_B_AI_Coding_Methodology_v1.0.docx")

    metadata = {
        "title": "LLM-Assisted Coding for Systematic Reviews: A Three-Model "
                 "Comparative Framework Applied to Educational Meta-Analysis",
        "title_short": "LLM-Assisted Coding",
        "running_head": "LLM-ASSISTED CODING FOR SYSTEMATIC REVIEWS",
        "author": "Hosung You",
        "affiliation": "College of Education, Pennsylvania State University",
        "author_note": [
            "Hosung You  https://orcid.org/[ORCID-ID]",
            "Correspondence concerning this article should be addressed to "
            "Hosung You, College of Education, Pennsylvania State University, "
            "University Park, PA 16802. Email: hosung@psu.edu",
            "Data Availability Statement: The complete dataset, AI prompts, "
            "extraction logs, and analysis code are available at [OSF Repository Link].",
            "Conflict of Interest: The author declares no conflicts of interest.",
            "Funding: This research received no external funding.",
        ],
    }

    doc = create_apa_document()
    add_running_head(doc, metadata["running_head"])
    add_page_number(doc)
    process_markdown_to_apa(md_path, doc, metadata)

    doc.save(out_path)
    size_kb = os.path.getsize(out_path) / 1024
    print(f"  Paper B: {out_path}")
    print(f"  Size: {size_kb:.1f} KB")
    return out_path


if __name__ == "__main__":
    print("Converting manuscripts to APA 7th Word format...")
    print()
    a = convert_paper_a()
    print()
    b = convert_paper_b()
    print()
    print("Done.")
