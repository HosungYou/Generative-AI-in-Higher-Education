#!/usr/bin/env python3
"""Generate the GenAI Meta-Analysis Coding Manual v10.0 as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants ──
NAVY = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = "D9E2F3"
DARK_BLUE_HEX = "1F4E79"
OUTPUT = "/Users/hosung/Generative-AI-in-Higher-Education/docs/GenAI_MetaAnalysis_Coding_Manual_v10.docx"

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size in [(1, 16), (2, 14), (3, 12)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = NAVY
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{DARK_BLUE_HEX}"/>')
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell._tc.get_or_add_tcPr().append(shading)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.paragraphs[0].runs else None
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_GRAY}"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    return t


def add_box(text, indent=0):
    """Add monospace formatted text block for workflows."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    return p


# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Coding Manual for Systematic Review\nand Meta-Analysis")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = NAVY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Generative AI Effects on Learning Outcomes\nin Higher Education")
run.font.size = Pt(18)
run.font.color.rgb = NAVY

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 10.0\nFebruary 2026")
run.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Corresponding Template: GenAI_MetaAnalysis_v10_TEMPLATE.xlsx")
run.font.size = Pt(11)
run.font.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Introduction and Purpose",
    "2. Eligibility Criteria (PICOS)",
    "3. Coding Workflow Overview",
    "4. Coder Training Protocol",
    "5. Study-Level Coding Instructions (STUDY_CHAR)",
    "6. Effect Size Coding Instructions (EFFECT_SIZES)",
    "7. Raw Extraction Protocol (RAW_EXTRACTION)",
    "8. Effect Size Calculation Protocol (ES_CALCULATIONS)",
    "9. AI Coding Protocol — Paper B (AI_CODING)",
    "10. Inter-Rater Reliability Protocol",
    "11. Discrepancy Resolution Protocol (DISCREPANCY_LOG)",
    "12. Study Exclusion Protocol (EXCLUSION_LOG)",
    "13. Database Search Protocol (SEARCH_LOG)",
    "14. Quality Assurance Checklist",
    "15. References",
    "Appendix A: Decision Trees",
    "Appendix B: Frequently Asked Questions",
    "Appendix C: Excel Template Sheet Descriptions",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction and Purpose", level=1)

doc.add_paragraph(
    "This coding manual provides standardized procedures for extracting, coding, and verifying "
    "data from primary studies included in a systematic review and meta-analysis examining the "
    "effects of generative artificial intelligence (GenAI) on learning outcomes in higher education. "
    "The manual accompanies the Excel coding template GenAI_MetaAnalysis_v10_TEMPLATE.xlsx, "
    "a 9-sheet workbook designed for transparent, reproducible data collection."
)

doc.add_heading("1.1 Two-Paper Strategy", level=2)
doc.add_paragraph(
    "This manual supports two complementary publications:"
)
add_bullet("Paper A (Meta-Analysis): Examines the overall and moderated effects of GenAI on higher "
           "education learning outcomes using a three-level Bayesian meta-analysis. Uses ONLY the "
           "human-verified gold standard dataset.")
add_bullet("Paper B (Methodology): Compares AI-assisted coding (Claude, GPT-4o, Gemini) against "
           "the human gold standard to evaluate LLM accuracy for systematic review data extraction. "
           "Uses both human and AI coding data.")

doc.add_heading("1.2 How to Use This Manual", level=2)
doc.add_paragraph(
    "Each coder should: (a) read this manual in full before beginning coding; "
    "(b) keep it open as a reference during coding; (c) consult the CODEBOOK sheet in the Excel "
    "template for variable definitions and valid values; and (d) record all ambiguous decisions in "
    "the DISCREPANCY_LOG sheet. When the manual does not cover a specific case, flag the item and "
    "discuss at the next consensus meeting. All new decisions are added to Appendix B (FAQ)."
)

doc.add_heading("1.3 Version History", level=2)
add_table(
    ["Version", "Date", "Changes"],
    [
        ["10.0", "2026-02-16", "Initial release — complete recoding from original PDFs"],
        ["—", "—", "Prior versions (v5–v9) archived in data/_legacy/"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 2: ELIGIBILITY CRITERIA
# ═══════════════════════════════════════════════════════════════
doc.add_heading("2. Eligibility Criteria", level=1)

doc.add_paragraph(
    "Eligibility is defined using the PICOS framework (Population, Intervention, Comparison, "
    "Outcomes, Study design), following PRISMA 2020 guidelines (Page et al., 2021)."
)

doc.add_heading("2.1 PICOS Framework", level=2)
add_table(
    ["Element", "Criterion", "Notes"],
    [
        ["Population", "Higher education students (undergraduate or graduate)", "Excludes K-12, professional training outside degree programs"],
        ["Intervention", "Generative AI tool (released after November 2022)", "ChatGPT, GPT-3.5/4/4o, Claude, Gemini, Copilot, Custom LLM, etc."],
        ["Comparison", "Control or comparison condition", "No AI, traditional instruction, alternative AI, waitlist"],
        ["Outcomes", "Quantitative learning outcomes", "Cognitive, affective, behavioral, or metacognitive with sufficient statistics"],
        ["Study design", "Experimental or quasi-experimental", "RCT or quasi-experimental with comparison group"],
    ]
)

doc.add_heading("2.2 Inclusion Criteria", level=2)
criteria = [
    "Higher education setting (undergraduate or graduate level)",
    "Generative AI tool as the primary intervention (released after November 2022)",
    "Control or comparison condition present (no AI, traditional, alternative AI, or waitlist)",
    "Quantitative learning outcomes reported with sufficient statistics for effect size calculation",
    "Experimental (RCT) or quasi-experimental design",
    "Published in English",
    "Peer-reviewed journal article, conference paper, or doctoral dissertation",
]
for c in criteria:
    add_bullet(c)

doc.add_heading("2.3 Exclusion Criteria", level=2)
exclusions = [
    "K-12 or non-degree professional training settings",
    "AI tools that are not generative (e.g., traditional adaptive learning systems, rule-based chatbots)",
    "No comparison/control group (single-group pre-post only)",
    "Qualitative-only outcomes or insufficient statistics for effect size calculation",
    "Review articles, meta-analyses, commentaries, or opinion pieces",
    "Non-English publications",
    "Preprints without peer review (unless subsequently published)",
]
for e in exclusions:
    add_bullet(e)

doc.add_heading("2.4 Borderline Decision Rules", level=2)
add_bullet("Mixed populations (e.g., undergrad + high school): Include ONLY if higher education "
           "results are reported separately or ≥80% of participants are in higher education.")
add_bullet("AI tools with both generative and non-generative features (e.g., AI tutoring system "
           "with ChatGPT integration): Include if the generative component is the primary intervention.")
add_bullet("Studies with multiple comparison groups: Include all pairwise comparisons against the "
           "GenAI condition as separate effect sizes.")
add_bullet("Dissertations subsequently published as journal articles: Include only the journal version "
           "to avoid duplicate data.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 3: CODING WORKFLOW
# ═══════════════════════════════════════════════════════════════
doc.add_heading("3. Coding Workflow Overview", level=1)

doc.add_paragraph(
    "The coding process follows a six-phase workflow designed to maximize reliability and "
    "transparency, consistent with Cochrane Handbook Chapter 5 recommendations for dual "
    "independent extraction (Higgins et al., 2023)."
)

doc.add_heading("3.1 Workflow Diagram", level=2)
workflow = """
┌─────────────────────────────────────────────────────┐
│  PHASE 1: SEARCH & SCREENING                       │
│  Database search → Deduplication → Title/Abstract   │
│  screening (dual) → Full-text screening (dual)      │
│  → Final study pool                                 │
└──────────────────────┬──────────────────────────────┘
                       ▼
┌─────────────────────────────────────────────────────┐
│  PHASE 2: INDEPENDENT DUAL CODING                   │
│  Coder A codes ALL studies independently             │
│  Coder B codes ALL studies independently             │
│  ★ No communication during this phase ★             │
└──────────────────────┬──────────────────────────────┘
                       ▼
┌─────────────────────────────────────────────────────┐
│  PHASE 3: RELIABILITY CHECK                         │
│  Calculate Cohen's κ (categorical variables)         │
│  Calculate ICC (continuous variables)                │
│  Document initial agreement rates                    │
└──────────────────────┬──────────────────────────────┘
                       ▼
┌─────────────────────────────────────────────────────┐
│  PHASE 4: CONSENSUS MEETING                         │
│  Compare all coded values side by side               │
│  Identify and classify all discrepancies             │
│  Discuss each disagreement with PDF open             │
│  Attempt to reach consensus                          │
└──────────────────────┬──────────────────────────────┘
                       ▼
┌─────────────────────────────────────────────────────┐
│  PHASE 5: THIRD REVIEWER  (if needed, ~5-15%)       │
│  Unresolved items → Third reviewer examines PDF      │
│  independently → Final binding decision              │
│  → Document rationale in DISCREPANCY_LOG             │
└──────────────────────┬──────────────────────────────┘
                       ▼
┌─────────────────────────────────────────────────────┐
│  PHASE 6: GOLD STANDARD FINALIZATION                │
│  Merge all consensus values into master dataset      │
│  Verify all effect size calculations                 │
│  Final quality assurance check → Lock dataset        │
└─────────────────────────────────────────────────────┘
"""
add_box(workflow)

doc.add_heading("3.2 Estimated Timeline", level=2)
add_table(
    ["Phase", "Duration", "Personnel"],
    [
        ["1. Search & Screening", "1–2 weeks", "Both coders + librarian"],
        ["2. Independent Dual Coding", "3–4 weeks", "Both coders (parallel)"],
        ["3. Reliability Check", "1–2 days", "Lead researcher"],
        ["4. Consensus Meetings", "1–2 weeks", "Both coders (joint)"],
        ["5. Third Reviewer", "As needed", "Third reviewer"],
        ["6. Finalization", "2–3 days", "Lead researcher"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 4: CODER TRAINING
# ═══════════════════════════════════════════════════════════════
doc.add_heading("4. Coder Training Protocol", level=1)

doc.add_paragraph(
    "Training follows a six-phase approach based on best practices from the Cochrane Handbook "
    "(Higgins et al., 2023) and Pigott and Polanin (2020). Both coders must complete all phases "
    "before proceeding to full coding."
)

add_table(
    ["Phase", "Duration", "Activities", "Success Criterion"],
    [
        ["1. Pre-Training", "1 week", "Read coding manual; review 3 sample papers", "Confirm readiness"],
        ["2. Training Session 1", "4 hours", "Project overview, manual walkthrough, guided coding of 2 studies, independent coding of 3 studies", "Discussion completed"],
        ["3. Pilot Coding", "1–2 weeks", "Independent dual coding of 10 studies", "IRR calculated"],
        ["4. Pilot Review", "2 hours", "Review IRR results, discuss all discrepancies, revise manual", "Manual updated"],
        ["5. Training Session 2", "2 hours", "Re-train on low-IRR variables, code 3 additional studies", "κ ≥ 0.70 achieved"],
        ["6. Reliability Gate", "—", "Verify minimum thresholds met before full coding", "κ ≥ 0.70, ICC ≥ 0.75"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Important: ")
run.bold = True
p.add_run("If reliability thresholds are not met after Phase 5, repeat Phases 3–5 with a new "
          "set of 10 studies until thresholds are achieved. Document all training iterations.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 5: STUDY-LEVEL CODING
# ═══════════════════════════════════════════════════════════════
doc.add_heading("5. Study-Level Coding Instructions", level=1)
doc.add_paragraph(
    "This section corresponds to Sheet 2: STUDY_CHAR in the Excel template. "
    "Each study receives exactly one row. All 34 variables are described below, grouped by domain. "
    "Each variable includes its operational definition and decision rules, following Pigott and Polanin's "
    "(2020) requirement for both theoretical and operational definitions."
)

# 5.1 Identification
doc.add_heading("5.1 Identification Variables (7)", level=2)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["study_id", "Integer", "Sequential unique identifier assigned during screening. Never reuse IDs of excluded studies.", "1, 2, 3, ..."],
        ["first_author", "String", "Family name of first author exactly as published.", "Kim"],
        ["year", "Integer", "Publication year. For online-first, use the year of first online availability.", "2025"],
        ["title", "String", "Full title as published. Do not abbreviate.", "Effects of ChatGPT on..."],
        ["doi", "String", "Digital Object Identifier. Include full URL format if no DOI (e.g., URL).", "10.1234/example"],
        ["source_type", "Categorical", "journal | conference | dissertation | preprint. Use 'journal' for online-first articles in journals.", "journal"],
        ["search_source", "String", "Database where study was first identified.", "ERIC"],
    ]
)

# 5.2 Study Design
doc.add_heading("5.2 Study Design Variables (5)", level=2)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["study_design", "Categorical", "RCT: random assignment explicitly stated. quasi_experimental: non-random comparison groups.", "RCT"],
        ["random_assignment", "Categorical", "yes: explicitly stated. no: comparison but not randomized. unclear: not described.", "yes"],
        ["matching_method", "Categorical", "none | propensity | stratified | other. Code 'none' if random assignment used.", "none"],
        ["blinding", "Categorical", "none | single | double | unclear. In education, most studies are 'none' or 'unclear'.", "none"],
        ["attrition_reported", "Categorical", "yes: dropout numbers reported. no: not mentioned.", "yes"],
    ]
)

# 5.3 Sample
doc.add_heading("5.3 Sample Variables (6)", level=2)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["n_total", "Integer", "Total analyzed participants (not enrolled). Use post-attrition N.", "120"],
        ["academic_level", "Categorical", "undergraduate | graduate | mixed. 'mixed' only if both levels included.", "undergraduate"],
        ["discipline", "Categorical", "CS | education | language | medicine | STEM_other | social_science | other", "language"],
        ["country", "String", "Country where study was conducted. Use ISO standard English name.", "South Korea"],
        ["institution_type", "Categorical", "public | private | unclear", "public"],
        ["prior_knowledge_controlled", "Categorical", "yes: pre-test or prior GPA used as covariate/matching. no: not controlled.", "yes"],
    ]
)

# 5.4 Intervention
doc.add_heading("5.4 Intervention Variables (7)", level=2)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["genai_tool", "Categorical", "ChatGPT | GPT-3.5 | GPT-4 | GPT-4o | Claude | Gemini | Copilot | Custom_LLM | Other | Not_specified. If study says 'ChatGPT' without version, code 'ChatGPT'.", "GPT-4"],
        ["genai_version_detail", "String", "Free text for specific version info (e.g., 'GPT-4-turbo-2024-01').", "GPT-4-turbo"],
        ["genai_role", "Categorical", "tutor | feedback_provider | writing_assistant | problem_solver | content_generator | other", "feedback_provider"],
        ["control_condition", "Categorical", "no_AI | traditional_instruction | alternative_AI | waitlist | other", "traditional_instruction"],
        ["duration_weeks", "Numeric", "Intervention duration in weeks. Convert days to weeks (÷7). Single session = 0.14.", "4"],
        ["total_sessions", "Integer", "Number of sessions. If not reported, leave blank.", "8"],
        ["implementation_context", "Categorical", "in_class | homework | both | lab | online", "in_class"],
    ]
)

# 5.5 RoB 2.0
doc.add_heading("5.5 Risk of Bias 2.0 Variables (7)", level=2)
doc.add_paragraph(
    "Assess risk of bias using the Cochrane Risk of Bias 2.0 tool (Sterne et al., 2019). "
    "Rate each domain as low | some_concerns | high based on signaling questions."
)
add_table(
    ["Variable", "Domain", "Key Question"],
    [
        ["rob_randomization", "Randomization process", "Was the allocation sequence random? Was allocation concealed?"],
        ["rob_deviations", "Deviations from intended interventions", "Were participants aware of assignment? Were there deviations?"],
        ["rob_missing_data", "Missing outcome data", "Were outcome data available for all or nearly all participants?"],
        ["rob_measurement", "Measurement of outcome", "Was the outcome measure appropriate? Were assessors blinded?"],
        ["rob_selection", "Selection of reported result", "Were multiple outcome measures used? Was there selective reporting?"],
        ["overall_rob", "Overall judgment", "low: all domains low. some_concerns: ≥1 some_concerns. high: ≥1 high."],
        ["rob_notes", "Free text", "Explain any 'high' or 'some_concerns' ratings briefly."],
    ]
)

# 5.6 Source Management
doc.add_heading("5.6 Source Management Variables (2)", level=2)
add_table(
    ["Variable", "Type", "Definition", "Example"],
    [
        ["human_coder", "String", "Initials of the coder who completed this row.", "HY"],
        ["coding_date", "Date", "Date this row was completed (YYYY-MM-DD).", "2026-03-15"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 6: EFFECT SIZE CODING
# ═══════════════════════════════════════════════════════════════
doc.add_heading("6. Effect Size Coding Instructions", level=1)
doc.add_paragraph(
    "This section corresponds to Sheet 3: EFFECT_SIZES. Each effect size receives one row. "
    "A single study may contribute multiple effect sizes (e.g., different outcome measures, "
    "different time points, different comparison groups)."
)

doc.add_heading("6.1 Complete Variable Reference (35 variables)", level=2)
doc.add_paragraph(
    "The following tables list every variable in the EFFECT_SIZES sheet with operational "
    "definitions and coding rules. Detailed guidance for complex variables follows in subsequent subsections."
)

doc.add_heading("Linking Variables (3)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["study_id", "Integer", "Must match study_id in STUDY_CHAR.", "1"],
        ["es_id", "String", "Format: S###_E## where ### = zero-padded study_id, ## = effect size number.", "S001_E01"],
        ["es_sequence", "Integer", "Ordinal position of this effect size within the study (1, 2, 3, ...).", "1"],
    ]
)

doc.add_heading("Sample — Effect Size Level (2)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["n_treatment", "Integer", "Number of participants analyzed in the treatment/GenAI group for THIS outcome. May differ from n_total if study has multiple treatment arms.", "45"],
        ["n_control", "Integer", "Number of participants analyzed in the control group for THIS outcome. If shared control, divide by number of comparisons.", "42"],
    ]
)

doc.add_heading("Outcome Classification (7)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["outcome_name", "String", "Descriptive label for the outcome as named in the study.", "Critical thinking score"],
        ["outcome_dimension", "Categorical", "cognitive | affective | behavioral | metacognitive. See decision tree in 6.2.", "cognitive"],
        ["blooms_level", "Categorical", "remember | understand | apply | analyze | evaluate | create. See guide in 6.3.", "analyze"],
        ["blooms_order", "Categorical", "lower (remember/understand/apply) | higher (analyze/evaluate/create).", "higher"],
        ["measurement_type", "Categorical", "test | questionnaire | rubric | log_data | think_aloud | mixed", "test"],
        ["measurement_timing", "Categorical", "immediate_posttest: within 1 week of intervention. delayed_posttest: >1 week after. during_intervention: measured while intervention ongoing.", "immediate_posttest"],
        ["reliability_alpha", "Numeric", "Cronbach's alpha or equivalent reliability coefficient for the measure. Leave blank if not reported.", "0.87"],
    ]
)

doc.add_heading("Raw Statistics — Post-test (4)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["m_treatment_post", "Numeric", "Post-test mean for the treatment/GenAI group.", "82.3"],
        ["sd_treatment_post", "Numeric", "Post-test SD for the treatment group. If SE reported, convert: SD = SE × √n.", "11.5"],
        ["m_control_post", "Numeric", "Post-test mean for the control group.", "75.1"],
        ["sd_control_post", "Numeric", "Post-test SD for the control group.", "12.8"],
    ]
)

doc.add_heading("Raw Statistics — Pre-test (4, optional)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["m_treatment_pre", "Numeric", "Pre-test mean for treatment group. Extract when available for sensitivity analysis.", "70.2"],
        ["sd_treatment_pre", "Numeric", "Pre-test SD for treatment group.", "10.1"],
        ["m_control_pre", "Numeric", "Pre-test mean for control group.", "69.8"],
        ["sd_control_pre", "Numeric", "Pre-test SD for control group.", "10.5"],
    ]
)

doc.add_heading("Alternative Statistics (4, when M/SD unavailable)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["reported_stat_type", "Categorical", "t | F | chi2 | p_only | d | eta_sq | r | OR | other. Record the type of statistic available.", "t"],
        ["reported_stat_value", "Numeric", "The numeric value of the reported statistic.", "2.45"],
        ["reported_df", "Numeric", "Degrees of freedom associated with the statistic. For t: df = n1+n2-2.", "86"],
        ["reported_p", "Numeric", "p-value as reported. Record exact value if given (e.g., 0.003), not 'p < .05'.", "0.016"],
    ]
)

doc.add_heading("Calculated Effect Sizes (6)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["hedges_g", "Numeric", "Hedges' g (bias-corrected standardized mean difference). Positive = GenAI group outperformed control.", "0.54"],
        ["se_g", "Numeric", "Standard error of Hedges' g.", "0.15"],
        ["var_g", "Numeric", "Variance of Hedges' g. var_g = se_g².", "0.0225"],
        ["ci_lower_95", "Numeric", "Lower bound of 95% confidence interval. ci_lower = g - 1.96 × se_g.", "0.25"],
        ["ci_upper_95", "Numeric", "Upper bound of 95% confidence interval. ci_upper = g + 1.96 × se_g.", "0.83"],
        ["calculation_method", "Categorical", "post_means | change_scores | t_to_g | F_to_g | p_to_g | reported_d_to_g | eta_to_g | r_to_g. Must match the actual formula used.", "post_means"],
    ]
)

doc.add_heading("Source Tracking (2)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["source_location", "String", "Exact location in the paper where key statistics were found.", "Table 3, p.12"],
        ["extraction_notes", "String", "Any notes about extraction difficulties, assumptions, or author contact.", "SD estimated from CI"],
    ]
)

doc.add_heading("Source Management (3)", level=3)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["human_coder", "String", "Initials of the coder who completed this row.", "HY"],
        ["coding_date", "Date", "Date this row was completed (YYYY-MM-DD).", "2026-03-15"],
        ["verification_status", "Categorical", "initial: first pass by one coder. double_coded: independently coded by both. consensus_resolved: final value after discrepancy resolution.", "double_coded"],
    ]
)

doc.add_heading("6.2 Effect Size Identification", level=2)
doc.add_paragraph(
    "Each effect size is identified by es_id in the format S###_E## "
    "(e.g., S001_E01, S001_E02). The es_sequence variable records the ordinal position "
    "within the study (1, 2, 3, ...)."
)

doc.add_heading("6.2 Outcome Classification", level=2)
doc.add_paragraph(
    "Classify each outcome using outcome_dimension. Use the following decision tree:"
)
decision_tree = """
  Is the outcome about knowledge, skills, or academic performance?
  ├─ YES → COGNITIVE
  │    Examples: test scores, grades, problem-solving accuracy,
  │    writing quality, code quality
  │
  └─ NO → Does it measure attitudes, emotions, or motivation?
       ├─ YES → AFFECTIVE
       │    Examples: satisfaction, self-efficacy, motivation,
       │    perceived usefulness, anxiety, engagement
       │
       └─ NO → Does it measure observable actions or behaviors?
            ├─ YES → BEHAVIORAL
            │    Examples: participation frequency, time-on-task,
            │    help-seeking, strategy use, collaboration
            │
            └─ NO → METACOGNITIVE
                 Examples: self-regulation, planning, monitoring,
                 reflection, metacognitive awareness
"""
add_box(decision_tree)

doc.add_heading("6.3 Bloom's Taxonomy Classification", level=2)
add_table(
    ["Level", "blooms_level", "blooms_order", "GenAI Study Examples"],
    [
        ["Remember", "remember", "lower", "Factual recall tests, multiple-choice knowledge quizzes"],
        ["Understand", "understand", "lower", "Comprehension tests, explanation tasks, summarization quality"],
        ["Apply", "apply", "lower", "Programming assignments, problem sets, case application"],
        ["Analyze", "analyze", "higher", "Critical thinking tests, data analysis tasks, compare-contrast essays"],
        ["Evaluate", "evaluate", "higher", "Peer review quality, argument evaluation, critique writing"],
        ["Create", "create", "higher", "Creative writing, project design, original code generation"],
    ]
)

doc.add_heading("6.4 Statistics Extraction Priority", level=2)
doc.add_paragraph(
    "Extract statistics following this hierarchy (Lipsey & Wilson, 2001). "
    "Always prefer higher-priority statistics when available. "
    "Record the calculation_method field to document which approach was used."
)
add_table(
    ["Priority", "Statistics Needed", "calculation_method", "Notes"],
    [
        ["1 (Best)", "M, SD, n for both groups (post-test)", "post_means", "Preferred: between-group post-test comparison"],
        ["2", "M_change, SD_change, n + pre-post r", "change_scores", "Only if post-test means not available AND r reported"],
        ["3", "t-value, n₁, n₂", "t_to_g", "Independent-samples t-test only (NOT paired)"],
        ["4", "F-value (1 df numerator), n₁, n₂", "F_to_g", "One-way ANOVA or equivalent"],
        ["5", "p-value with direction, n₁, n₂", "p_to_g", "Last resort — least precise"],
        ["6", "Reported Cohen's d", "reported_d_to_g", "Apply Hedges' correction J"],
        ["—", "η², n₁, n₂", "eta_to_g", "Convert η² to d, then correct"],
        ["—", "r, n", "r_to_g", "Convert correlation to d, then correct"],
    ]
)

p = doc.add_paragraph()
run = p.add_run("AVOID: ")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run("Single-group pre-post effect sizes without a control group should NOT be computed. "
          "Pre-post designs conflate intervention effects with natural change and regression to "
          "the mean (Cuijpers et al., 2017).")

doc.add_heading("6.5 Where to Find Statistics", level=2)
add_bullet("Tables: Most common location. Check all tables, not just the first results table.")
add_bullet("Results section text: Sometimes means and SDs appear inline.")
add_bullet("Supplementary materials: Check online appendices and supplementary files.")
add_bullet("Figures: As a last resort, use WebPlotDigitizer to extract values from bar charts. "
           "Mark extraction_confidence as 'low' in RAW_EXTRACTION.")
add_bullet("If insufficient data: Contact corresponding author (document attempt in extraction_notes).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 7: RAW EXTRACTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading("7. Raw Extraction Protocol", level=1)
doc.add_paragraph(
    "This section corresponds to Sheet 4: RAW_EXTRACTION. Every numeric value extracted from "
    "a primary study must be documented with its exact source location. This creates an audit "
    "trail from the published paper to the coded dataset."
)

doc.add_heading("7.1 Complete Variable Reference (12 variables)", level=2)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["study_id", "Integer", "Must match study_id in STUDY_CHAR and EFFECT_SIZES.", "1"],
        ["es_id", "String", "Must match es_id in EFFECT_SIZES. Links this extraction to a specific effect size.", "S001_E01"],
        ["extraction_id", "String", "Unique ID for each extracted value. Format: S###_E##_X## (e.g., S001_E01_X01). Sequential within each effect size.", "S001_E01_X01"],
        ["page_number", "Integer", "PDF page number where the value appears.", "12"],
        ["table_or_figure", "String", "Specific table, figure, or text location (e.g., 'Table 3', 'Figure 2', 'text', 'supplementary'). Free text — no restricted values.", "Table 3"],
        ["section", "Categorical", "results | discussion | appendix. Section of the paper where value was found.", "results"],
        ["exact_quote", "String", "Copy the sentence or cell containing the value verbatim. Include surrounding context for clarity.", "M = 82.3, SD = 11.5, n = 45"],
        ["variable_extracted", "String", "Name of the EFFECT_SIZES variable this value maps to (e.g., 'm_treatment_post', 'sd_control_post').", "m_treatment_post"],
        ["value", "String", "The extracted numeric or text value exactly as reported.", "82.3"],
        ["unit_or_scale", "String", "Unit of measurement or scale range (e.g., '0-100', 'Likert 1-5', 'seconds').", "0-100"],
        ["extraction_confidence", "Categorical", "high | medium | low. See confidence definitions in 7.2.", "high"],
        ["confidence_notes", "String", "Explanation if confidence is medium or low. Blank if high.", ""],
    ]
)

doc.add_heading("7.2 Documentation Guidelines", level=2)
add_bullet("Create one row per extracted statistic. A single effect size typically has 4–8 rows (M, SD, n for each group).")
add_bullet("variable_extracted must exactly match a column name in EFFECT_SIZES (e.g., 'm_treatment_post', not 'treatment mean').")
add_bullet('exact_quote: Copy the sentence or table cell containing the value. '
           'Example: "The treatment group scored higher (M = 82.3, SD = 11.5, n = 45)."')

doc.add_heading("7.2 Confidence Ratings", level=2)
add_table(
    ["Rating", "Definition", "Action Required"],
    [
        ["high", "Value is clearly and unambiguously stated in the source.", "None — proceed with coding."],
        ["medium", "Value requires minor inference (e.g., calculated from CI, read from figure).", "Document inference method in confidence_notes."],
        ["low", "Value requires substantial inference or assumption.", "Flag for consensus discussion. Document all assumptions."],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 8: EFFECT SIZE CALCULATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading("8. Effect Size Calculation Protocol", level=1)
doc.add_paragraph(
    "This section corresponds to Sheet 5: ES_CALCULATIONS. Each calculated effect size must "
    "be documented step-by-step to ensure reproducibility. All formulas follow Lipsey and Wilson (2001)."
)

doc.add_heading("8.1 Primary Method: Post-Test Means", level=2)
formulas_1 = """
  Step 1: Calculate pooled SD
    SD_pooled = sqrt[((n1-1)*SD1^2 + (n2-1)*SD2^2) / (n1+n2-2)]

  Step 2: Calculate Cohen's d
    d = (M_treatment - M_control) / SD_pooled

  Step 3: Calculate Hedges' correction factor J
    J = 1 - 3 / (4*(n1+n2-2) - 1)

  Step 4: Calculate Hedges' g
    g = J * d

  Step 5: Calculate standard error of g
    SE(g) = sqrt[ (n1+n2)/(n1*n2) + g^2/(2*(n1+n2)) ]

  Step 6: Calculate variance and 95% CI
    Var(g) = SE(g)^2
    CI_lower = g - 1.96 * SE(g)
    CI_upper = g + 1.96 * SE(g)
"""
add_box(formulas_1)

doc.add_heading("8.2 Alternative: t-Value Conversion", level=2)
formulas_2 = """
  d = t * sqrt(1/n1 + 1/n2)
  Then apply Steps 3-6 above.
"""
add_box(formulas_2)

doc.add_heading("8.3 Alternative: F-Value Conversion", level=2)
formulas_3 = """
  d = sqrt( F * (1/n1 + 1/n2) )    [only for F with 1 df numerator]
  Then apply Steps 3-6 above.
"""
add_box(formulas_3)

doc.add_heading("8.4 Verification", level=2)
doc.add_paragraph(
    "After calculation, compare the hedges_g_final in ES_CALCULATIONS with the hedges_g in "
    "EFFECT_SIZES. If they do not match, set matches_sheet3 to 'no' and investigate."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 9: AI CODING (Paper B)
# ═══════════════════════════════════════════════════════════════
doc.add_heading("9. AI Coding Protocol — Paper B Only", level=1)
doc.add_paragraph(
    "This section corresponds to Sheet 6: AI_CODING. It describes the protocol for comparing "
    "three large language models against the human gold standard. This data is used exclusively "
    "in Paper B and does NOT affect the gold standard dataset used in Paper A."
)

doc.add_heading("9.1 Models", level=2)
add_table(
    ["Model", "Provider", "Expected Accuracy*"],
    [
        ["Claude (Sonnet/Opus)", "Anthropic", "72–80% (education domain)"],
        ["GPT-4o", "OpenAI", "72–80% (education domain)"],
        ["Gemini (Pro)", "Google", "72–80% (education domain)"],
    ]
)
doc.add_paragraph(
    "*Based on Goel et al. (2024) feasibility study reporting 72% accuracy in social sciences."
).italic = True

doc.add_heading("9.2 Prompt Design", level=2)
add_bullet("System prompt: Define role as 'meta-analysis data extraction assistant'.")
add_bullet("Include relevant variable definitions from CODEBOOK sheet.")
add_bullet("Provide the full study PDF as input.")
add_bullet("Request structured JSON output matching EFFECT_SIZES column names.")
add_bullet("Ask for confidence score (0–1) for each extracted value.")
add_bullet("Record the exact prompt version used (claude_prompt_version, etc.).")

doc.add_heading("9.3 Execution Rules", level=2)
add_bullet("Use the SAME prompt template for all three models (only system prompt preamble differs).")
add_bullet("Run each model independently — no chaining or using one model's output as input for another.")
add_bullet("Record raw output verbatim (truncated if >500 chars) in *_raw_output fields.")
add_bullet("Do NOT correct AI outputs. Preserve raw values for accuracy comparison.")

doc.add_heading("9.4 Consensus Methods (to be tested)", level=2)
add_bullet("Majority vote: Value agreed upon by ≥2 of 3 models.")
add_bullet("Weighted mean: Average weighted by each model's confidence score.")
add_bullet("Simple mean: Unweighted average (numeric fields only).")

doc.add_heading("9.5 Comparison Metrics", level=2)
add_bullet("Exact match rate: Proportion where AI value = human gold standard (per field, per model).")
add_bullet("Cohen's kappa: For categorical fields (study_design, outcome_dimension, blooms_level).")
add_bullet("ICC(2,1): For continuous fields (means, SDs, effect sizes).")
add_bullet("Mean Absolute Error: For numeric fields.")
add_bullet("Error taxonomy: misread | misinterpret | hallucinate | omit (classified by human reviewers).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 10: IRR PROTOCOL
# ═══════════════════════════════════════════════════════════════
doc.add_heading("10. Inter-Rater Reliability Protocol", level=1)
doc.add_paragraph(
    "Inter-rater reliability (IRR) is calculated BEFORE consensus resolution to reflect true "
    "initial agreement between coders. Only approximately one-third of published meta-analyses "
    "report IRR estimates (Pigott & Polanin, 2020); this protocol exceeds that standard by "
    "reporting comprehensive reliability metrics."
)

doc.add_heading("10.1 When to Calculate", level=2)
add_bullet("After independent dual coding is complete (Phase 3)")
add_bullet("BEFORE any consensus discussion (Phase 4)")
add_bullet("Interim checks: after every 20–25 studies during full coding")

doc.add_heading("10.2 Metrics and Thresholds", level=2)
add_table(
    ["Metric", "Applied To", "Minimum", "Target", "If Below Minimum"],
    [
        ["Cohen's κ", "Categorical variables", "0.67", "0.80", "Retrain + recode batch"],
        ["ICC(2,1)", "Continuous variables", "0.75", "0.90", "Retrain + recode batch"],
        ["% Agreement", "All variables (supplementary)", "80%", "90%", "Review variable definitions"],
    ]
)
doc.add_paragraph(
    "Thresholds based on Krippendorff's conservative standards (κ ≥ 0.67 for tentative conclusions, "
    "≥ 0.80 for definitive; cited in Hallgren, 2012) and Cicchetti's (1994) ICC guidelines "
    "(≥ 0.75 = excellent)."
)

doc.add_heading("10.3 Reporting Format", level=2)
doc.add_paragraph("Report IRR in a table structured as follows:")
add_table(
    ["Variable Category", "Metric", "Value", "95% CI", "n items"],
    [
        ["Study design (5 vars)", "Cohen's κ", "[value]", "[lower, upper]", "[k]"],
        ["Sample characteristics (6 vars)", "Cohen's κ", "[value]", "[lower, upper]", "[k]"],
        ["Intervention (7 vars)", "Cohen's κ", "[value]", "[lower, upper]", "[k]"],
        ["RoB 2.0 (6 vars)", "Cohen's κ", "[value]", "[lower, upper]", "[k]"],
        ["Outcome classification (4 vars)", "Cohen's κ", "[value]", "[lower, upper]", "[k]"],
        ["Post-test statistics (4 vars)", "ICC(2,1)", "[value]", "[lower, upper]", "[k]"],
        ["Calculated effect sizes (3 vars)", "ICC(2,1)", "[value]", "[lower, upper]", "[k]"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 11: DISCREPANCY RESOLUTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading("11. Discrepancy Resolution Protocol", level=1)
doc.add_paragraph(
    "This section corresponds to Sheet 7: DISCREPANCY_LOG. A clear, explicit process ensures "
    "every disagreement is resolved transparently and documented for reproducibility."
)

doc.add_heading("11.1 Discrepancy Classification", level=2)
add_table(
    ["Classification", "Definition", "Examples", "Resolution"],
    [
        ["MINOR", "Trivial differences unlikely to affect analysis", "Typos, rounding (≤0.01 for numeric), formatting differences", "Auto-resolve: take mean (numeric) or Coder A's value (categorical)"],
        ["MAJOR", "Substantive differences that could affect results", "Different category assigned, different values extracted, missing vs. present", "Must discuss at consensus meeting"],
    ]
)

doc.add_heading("11.2 Resolution Workflow", level=2)
resolution_flow = """
  STEP 1: AUTOMATED COMPARISON
  ├─ Export both coders' data side-by-side
  ├─ Flag ALL cells where values differ
  └─ Count: ___ minor, ___ major discrepancies

  STEP 2: AUTO-RESOLVE MINOR DISCREPANCIES
  ├─ Numeric (≤0.01 difference): take mean
  ├─ Categorical (typo/case): standardize
  └─ Log resolution in DISCREPANCY_LOG

  STEP 3: SCHEDULE CONSENSUS MEETING (within 48 hours)
  ├─ Both coders attend with PDFs open
  └─ Work through major discrepancies one by one

  STEP 4: FOR EACH MAJOR DISCREPANCY:
  ├─ Both coders re-read relevant section of PDF
  ├─ Each coder explains their reasoning (2 min each)
  ├─ Consult coding manual decision rules
  ├─ Check: Is there a clear right answer?
  │   ├─ YES → Adopt correct value, log error type
  │   └─ NO  → Discuss until consensus (max 15 min)
  │
  ├─ CONSENSUS REACHED?
  │   ├─ YES → Record final_value + rationale
  │   └─ NO  → ESCALATE TO STEP 5
  └─ Log everything in DISCREPANCY_LOG

  STEP 5: THIRD REVIEWER ARBITRATION
  ├─ Third reviewer receives:
  │   ├─ Original PDF
  │   ├─ Both coders' values
  │   └─ Discussion notes
  ├─ Reviews independently (no prior discussion)
  ├─ Makes FINAL BINDING decision
  └─ Documents rationale in DISCREPANCY_LOG
"""
add_box(resolution_flow)

doc.add_heading("11.3 Documentation in DISCREPANCY_LOG", level=2)
doc.add_paragraph("Every major discrepancy is recorded with these fields:")
add_table(
    ["Field", "Description"],
    [
        ["study_id, es_id", "Links to the specific study and effect size"],
        ["variable_name", "Which variable was in disagreement"],
        ["coder1_value", "Coder A's original value"],
        ["coder2_value", "Coder B's original value"],
        ["discrepancy_type", "human_human | ai_human | ai_ai"],
        ["magnitude", "minor | major"],
        ["final_value", "Consensus or third-reviewer value"],
        ["resolution_method", "discussion | third_reviewer | source_recheck | calculation_error"],
        ["resolved_by", "Initials of person(s) who resolved"],
        ["resolution_date", "Date resolved (YYYY-MM-DD)"],
    ]
)

doc.add_heading("11.4 Learning from Disagreements", level=2)
add_bullet("After every 20 studies, review discrepancy patterns.")
add_bullet("If a variable causes >25% disagreements, revise its definition in this manual.")
add_bullet("Add new decision rules to Appendix B (FAQ) based on recurring ambiguities.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 12: EXCLUSION LOG
# ═══════════════════════════════════════════════════════════════
doc.add_heading("12. Study Exclusion Protocol", level=1)
doc.add_paragraph(
    "This section corresponds to Sheet 8: EXCLUSION_LOG. Every study examined at full-text "
    "screening or later that is NOT included in the final dataset must be documented here "
    "with the reason for exclusion. This is required for the PRISMA flow diagram."
)

doc.add_heading("12.1 Complete Variable Reference (10 variables)", level=2)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["study_id", "Integer", "Sequential ID assigned during screening. Excluded studies retain their IDs (never reused).", "15"],
        ["first_author", "String", "Family name of first author.", "Zhang"],
        ["year", "Integer", "Publication year.", "2024"],
        ["title", "String", "Full title as published.", "AI Tutoring in K-12..."],
        ["exclusion_stage", "Categorical", "title_abstract: excluded during title/abstract screening. full_text: excluded after full-text review. post_coding: excluded after data extraction revealed problems.", "full_text"],
        ["exclusion_reason", "Categorical", "wrong_population | no_control | non_GenAI | insufficient_data | duplicate | not_peer_reviewed | wrong_design | other", "wrong_population"],
        ["detailed_rationale", "String", "Specific explanation. Be precise enough for another reviewer to verify.", "Study conducted in high school (Grade 11), not higher education."],
        ["screener1_decision", "Categorical", "include | exclude | uncertain. Screener 1's independent decision.", "exclude"],
        ["screener2_decision", "Categorical", "include | exclude | uncertain. Screener 2's independent decision.", "exclude"],
        ["final_decision", "Categorical", "include | exclude. Final consensus decision after discussion.", "exclude"],
    ]
)

doc.add_heading("12.2 When to Log", level=2)
add_bullet("Log ALL studies excluded at full-text screening (not title/abstract rejects unless sample tracked).")
add_bullet("Log studies excluded post-coding (e.g., discovered insufficient data during extraction).")
add_bullet("If screeners disagree (one 'include', one 'exclude'), record both decisions and resolve via discussion or third reviewer.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 13: SEARCH LOG
# ═══════════════════════════════════════════════════════════════
doc.add_heading("13. Database Search Protocol", level=1)
doc.add_paragraph(
    "This section corresponds to Sheet 9: SEARCH_LOG. Every database search must be "
    "documented for reproducibility, following PRISMA 2020 requirements (Page et al., 2021)."
)

doc.add_heading("13.1 Complete Variable Reference (5 variables)", level=2)
add_table(
    ["Variable", "Type", "Definition & Rules", "Example"],
    [
        ["database", "String", "Name of the database searched (e.g., PsycINFO, ERIC, Education Source, ProQuest, Semantic Scholar, OpenAlex, Web of Science).", "ERIC"],
        ["search_date", "Date", "Date the search was conducted (YYYY-MM-DD). Re-searches get new rows.", "2026-02-10"],
        ["search_string", "String", "Complete search query exactly as entered, including Boolean operators and field tags.", '("generative AI" OR "ChatGPT" OR "GPT-4") AND "higher education"'],
        ["results_count", "Integer", "Number of records returned by this search.", "342"],
        ["notes", "String", "Any relevant notes (e.g., filters applied, date range, search limitations).", "Limited to 2023-2026, English only"],
    ]
)

doc.add_heading("13.2 Search Documentation Requirements", level=2)
add_bullet("Record one row per database per search date. If you search ERIC twice on different dates, create two rows.")
add_bullet("Copy the exact search string — do not paraphrase or simplify.")
add_bullet("Document any database-specific syntax (e.g., MeSH terms, ERIC descriptors).")
add_bullet("If a database requires multiple searches (e.g., different field combinations), record each as a separate row.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 14: QA CHECKLIST
# ═══════════════════════════════════════════════════════════════
doc.add_heading("14. Quality Assurance Checklist", level=1)
doc.add_paragraph("Complete this checklist for EVERY study before finalizing:")

qa_items = [
    "All required fields in STUDY_CHAR are complete (no blank cells for required variables)",
    "study_id in STUDY_CHAR matches study_id in EFFECT_SIZES",
    "Every effect size has source_location documented (page/table reference)",
    "Hedges' g in EFFECT_SIZES matches recalculation in ES_CALCULATIONS",
    "All categorical values match the valid_values listed in CODEBOOK",
    "Sample sizes (n_treatment + n_control) are consistent across effect sizes within a study",
    "RoB 2.0 assessment is complete with rob_notes for any 'high' or 'some_concerns' rating",
    "Pre-test data extracted where available (even if primary analysis uses post-test only)",
    "calculation_method correctly reflects the actual formula used",
    "RAW_EXTRACTION has entries for all effect size statistics with page references",
]
for i, item in enumerate(qa_items, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"☐ {i}. ")
    run.bold = True
    p.add_run(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 13: REFERENCES
# ═══════════════════════════════════════════════════════════════
doc.add_heading("15. References", level=1)

refs = [
    "Cicchetti, D. V. (1994). Guidelines, criteria, and rules of thumb for evaluating normed and standardized assessment instruments in psychology. Psychological Assessment, 6(4), 284–290.",
    "Cooper, H. (2017). Research synthesis and meta-analysis: A step-by-step approach (5th ed.). Sage.",
    "Cuijpers, P., Weitz, E., Cristea, I. A., & Twisk, J. (2017). Pre-post effect sizes should be avoided in meta-analyses. Epidemiology and Psychiatric Sciences, 26(4), 364–368.",
    "Goel, A., et al. (2024). Exploring the use of a large language model for data extraction in systematic reviews: A rapid feasibility study. arXiv:2405.14445.",
    "Hallgren, K. A. (2012). Computing inter-rater reliability for observational data: An overview and tutorial. Tutorials in Quantitative Methods for Psychology, 8(1), 23–34.",
    "Higgins, J. P. T., Thomas, J., Chandler, J., Cumpston, M., Li, T., Page, M. J., & Welch, V. A. (Eds.). (2023). Cochrane handbook for systematic reviews of interventions (Version 6.4). Cochrane. www.training.cochrane.org/handbook",
    "Landis, J. R., & Koch, G. G. (1977). The measurement of observer agreement for categorical data. Biometrics, 33(1), 159–174.",
    "Lipsey, M. W., & Wilson, D. B. (2001). Practical meta-analysis. Sage.",
    "Page, M. J., McKenzie, J. E., Bossuyt, P. M., Boutron, I., Hoffmann, T. C., Mulrow, C. D., ... & Moher, D. (2021). The PRISMA 2020 statement: An updated guideline for reporting systematic reviews. BMJ, 372, n71.",
    "Pigott, T. D., & Polanin, J. R. (2020). Methodological guidance paper: High-quality meta-analysis in a systematic review. Review of Educational Research, 90(1), 24–46.",
    "Sterne, J. A. C., Savović, J., Page, M. J., Elbers, R. G., Blencowe, N. S., Boutron, I., ... & Higgins, J. P. T. (2019). RoB 2: A revised tool for assessing risk of bias in randomised trials. BMJ, 366, l4898.",
    "Waffenschmidt, S., Knelangen, M., Sieben, W., Bühn, S., & Pieper, D. (2019). Single screening versus conventional double screening for study selection in systematic reviews. BMC Medical Research Methodology, 19(1), 132.",
    "Wang, Y., et al. (2025). Artificial intelligence–assisted data extraction with a large language model: A study within reviews. Annals of Internal Medicine.",
    "Xie, Y., et al. (2025). Collaborative large language models for automated data extraction in living systematic reviews. Journal of Clinical Epidemiology.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDIX A: DECISION TREES
# ═══════════════════════════════════════════════════════════════
doc.add_heading("Appendix A: Decision Trees", level=1)

doc.add_heading("A.1 Inclusion/Exclusion Decision", level=2)
tree_a1 = """
  Is it in a higher education setting?
  ├─ NO  → EXCLUDE (wrong_population)
  └─ YES → Does it use a generative AI tool (post-Nov 2022)?
       ├─ NO  → EXCLUDE (non_GenAI)
       └─ YES → Is there a control/comparison group?
            ├─ NO  → EXCLUDE (no_control)
            └─ YES → Are quantitative outcomes reported?
                 ├─ NO  → EXCLUDE (insufficient_data)
                 └─ YES → Can an effect size be calculated?
                      ├─ NO  → EXCLUDE (insufficient_data)
                      └─ YES → Is it experimental/quasi-experimental?
                           ├─ NO  → EXCLUDE (wrong_design)
                           └─ YES → INCLUDE
"""
add_box(tree_a1)

doc.add_heading("A.2 GenAI Tool Classification", level=2)
tree_a2 = """
  Does the paper name a specific model?
  ├─ NO  → Code as 'Not_specified'
  └─ YES → Is it an OpenAI model?
       ├─ YES → Which version?
       │    ├─ "ChatGPT" (no version)  → ChatGPT
       │    ├─ "GPT-3.5" / "GPT-3.5-turbo" → GPT-3.5
       │    ├─ "GPT-4" / "GPT-4-turbo" → GPT-4
       │    └─ "GPT-4o" → GPT-4o
       └─ NO → Is it Anthropic Claude? → Claude
            └─ Is it Google Gemini/Bard? → Gemini
                 └─ Is it GitHub Copilot? → Copilot
                      └─ Is it a custom/fine-tuned LLM? → Custom_LLM
                           └─ Otherwise → Other
"""
add_box(tree_a2)

doc.add_heading("A.3 Outcome Dimension Classification", level=2)
tree_a3 = """
  Knowledge / skills / performance test?
  ├─ YES → cognitive
  └─ NO → Attitude / motivation / emotion / satisfaction?
       ├─ YES → affective
       └─ NO → Observable behavior / participation / strategy use?
            ├─ YES → behavioral
            └─ NO → Self-regulation / metacognitive awareness / reflection?
                 ├─ YES → metacognitive
                 └─ Unclear → Discuss at consensus meeting
"""
add_box(tree_a3)

doc.add_heading("A.4 Bloom's Taxonomy Level", level=2)
tree_a4 = """
  Does the assessment require producing something new?
  ├─ YES → create (higher)
  └─ NO → Does it require making judgments?
       ├─ YES → evaluate (higher)
       └─ NO → Does it require breaking down into parts?
            ├─ YES → analyze (higher)
            └─ NO → Does it require using knowledge in new situations?
                 ├─ YES → apply (lower)
                 └─ NO → Does it require explaining concepts?
                      ├─ YES → understand (lower)
                      └─ NO → remember (lower)
"""
add_box(tree_a4)

doc.add_heading("A.5 Effect Size Calculation Method Selection", level=2)
tree_a5 = """
  Are M, SD, and n available for both groups (post-test)?
  ├─ YES → post_means
  └─ NO → Are change scores with pre-post r available?
       ├─ YES → change_scores
       └─ NO → Is a between-group t-value reported?
            ├─ YES → t_to_g
            └─ NO → Is an F-value (1 df numerator) reported?
                 ├─ YES → F_to_g
                 └─ NO → Is Cohen's d reported directly?
                      ├─ YES → reported_d_to_g
                      └─ NO → Is η² reported?
                           ├─ YES → eta_to_g
                           └─ NO → Is a p-value with direction available?
                                ├─ YES → p_to_g
                                └─ NO → Contact authors or EXCLUDE
"""
add_box(tree_a5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDIX B: FAQ
# ═══════════════════════════════════════════════════════════════
doc.add_heading("Appendix B: Frequently Asked Questions", level=1)

faqs = [
    ("Q1: What if a study uses multiple AI tools (e.g., ChatGPT + Copilot)?",
     "Code the PRIMARY tool used in the main intervention. If both are used equally, code the more prominent one in genai_tool and note the second in genai_version_detail."),
    ("Q2: What if sample sizes differ between pre-test and post-test?",
     "Use post-test (analyzed) sample sizes for n_treatment and n_control. Document the attrition in STUDY_CHAR (attrition_reported = yes) and note original N in extraction_notes."),
    ("Q3: How to handle studies with more than 2 groups?",
     "Create separate effect sizes for each pairwise comparison against the GenAI condition. For example, GenAI vs. No-AI = one ES, GenAI vs. Traditional = another ES. Adjust n for shared control groups (divide n_control by number of comparisons)."),
    ("Q4: What if SD is not reported but SE or CI is?",
     "Convert: SD = SE × √n. From 95% CI: SD = √n × (CI_upper - CI_lower) / 3.92. Document the conversion in calculation_method notes."),
    ("Q5: How to code custom or fine-tuned LLMs?",
     "If the LLM is based on a known model (e.g., fine-tuned GPT-3.5), code the base model in genai_tool and note 'fine-tuned' in genai_version_detail. If entirely custom, code as Custom_LLM."),
    ("Q6: What counts as 'generative AI' vs. traditional AI/ML?",
     "Generative AI must be capable of generating novel text, code, or content (e.g., ChatGPT, Claude, Gemini, Copilot). Traditional AI (e.g., ITS, adaptive learning algorithms, recommender systems) that only selects/retrieves content = EXCLUDE."),
    ("Q7: What if a study reports only qualitative outcomes alongside quantitative ones?",
     "Code only the quantitative outcomes. Ignore qualitative findings for effect size extraction. Note 'mixed methods study' in extraction_notes."),
    ("Q8: What if the same study appears as both conference paper and journal article?",
     "Include ONLY the journal article (most complete version). If the conference paper has unique data not in the journal version, note this in extraction_notes but still use only the journal version."),
    ("Q9: How to handle cluster randomization?",
     "If the study reports cluster-adjusted statistics (e.g., multilevel model results), use those. If not adjusted, note 'cluster randomization not adjusted' in rob_notes and consider sensitivity analysis excluding these studies."),
    ("Q10: What if the outcome is binary (pass/fail) rather than continuous?",
     "Convert odds ratio or risk ratio to Cohen's d using the Hasselblad-Hedges formula: d = ln(OR) × √3 / π. Code calculation_method as 'OR_to_g'. Note this conversion in extraction_notes."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)
    doc.add_paragraph()  # spacing

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDIX C: EXCEL SHEETS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("Appendix C: Excel Template Sheet Descriptions", level=1)

sheets = [
    ("CODEBOOK", "Contains all 144 variable definitions across the workbook. Each row describes one variable with its name, parent sheet, data type, valid values, coding rules, and examples. This is the master reference — consult it whenever unsure about how to code a variable."),
    ("STUDY_CHAR", "Study-level characteristics: one row per included study (34 variables). Covers identification, study design, sample demographics, intervention details, RoB 2.0 assessment, and source management. Complete this sheet fully for each study before moving to effect sizes."),
    ("EFFECT_SIZES", "Core analysis data: one row per effect size (35 variables). Multiple rows per study are expected when studies report multiple outcomes. This sheet is the primary input for both the three-level meta-analysis (Paper A) and the AI comparison (Paper B)."),
    ("RAW_EXTRACTION", "Audit trail: documents the exact source (page, table, quote) for every numeric value extracted. One row per extracted statistic — a single effect size may have 4–8 rows here (means, SDs, sample sizes). Essential for verification."),
    ("ES_CALCULATIONS", "Calculation documentation: shows the step-by-step derivation of each Hedges' g value. One row per effect size. Includes input values, formula used, intermediate results, and a verification flag (matches_sheet3). Ensures reproducibility."),
    ("AI_CODING", "Paper B data: records LLM extraction results for each study × variable combination. Contains columns for Claude, GPT-4o, and Gemini values alongside the human gold standard. Includes confidence scores, raw outputs, and consensus computations."),
    ("DISCREPANCY_LOG", "Resolution record: one row per disagreement between coders (human–human, AI–human, or AI–AI). Documents both initial values, the resolution method, final value, and who resolved it. Critical for transparency reporting."),
    ("EXCLUSION_LOG", "Screening record: one row per excluded study. Documents the exclusion stage, reason, and both screeners' independent decisions. Required for PRISMA flow diagram."),
    ("SEARCH_LOG", "Database search record: documents each database searched, the date, exact search string used, and number of results. Required for PRISMA reporting and reproducibility."),
]
for name, desc in sheets:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# HEADERS & FOOTERS
# ═══════════════════════════════════════════════════════════════
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "GenAI Meta-Analysis Coding Manual v10.0"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.style.font.size = Pt(9)
    hp.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Footer - page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

# ── Save ──
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"✓ Saved: {OUTPUT}")
print(f"✓ Size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
